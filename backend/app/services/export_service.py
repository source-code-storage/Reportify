"""
Export service for generating PDF and DOCX files from reports
"""

import io
import os
from typing import Optional, Dict, Any
from datetime import datetime


class ExportService:
    """Service for exporting reports to PDF and DOCX formats"""

    def __init__(self):
        """Initialize export service"""
        self._reportlab_available = None
        self._docx_available = None

    @property
    def reportlab_available(self) -> bool:
        """Check if ReportLab is available"""
        if self._reportlab_available is None:
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.platypus import SimpleDocTemplate

                self._reportlab_available = True
            except ImportError:
                self._reportlab_available = False
        return self._reportlab_available

    @property
    def docx_available(self) -> bool:
        """Check if python-docx is available"""
        if self._docx_available is None:
            try:
                from docx import Document

                self._docx_available = True
            except ImportError:
                self._docx_available = False
        return self._docx_available

    def export_to_pdf(
        self,
        report_title: str,
        sections: list,
        author: str = "",
        metadata: Dict[str, Any] = None,
    ) -> bytes:
        """
        Export report to PDF format.

        Args:
            report_title: Title of the report
            sections: List of sections with title and content
            author: Author name
            metadata: Additional metadata

        Returns:
            PDF file as bytes
        """
        if not self.reportlab_available:
            raise ImportError(
                "reportlab is required for PDF export. Install with: pip install reportlab"
            )

        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer,
            PageBreak,
            Table,
            TableStyle,
        )
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

        # Create buffer
        buffer = io.BytesIO()

        # Create document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72,
        )

        # Get styles
        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Heading1"],
            fontSize=24,
            spaceAfter=30,
            alignment=TA_CENTER,
        )

        heading_style = ParagraphStyle(
            "CustomHeading",
            parent=styles["Heading2"],
            fontSize=14,
            spaceBefore=20,
            spaceAfter=10,
        )

        body_style = ParagraphStyle(
            "CustomBody",
            parent=styles["Normal"],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceAfter=12,
        )

        # Build content
        story = []

        # Title page
        story.append(Spacer(1, 2 * inch))
        story.append(Paragraph(report_title, title_style))
        story.append(Spacer(1, 0.5 * inch))

        if author:
            author_style = ParagraphStyle(
                "Author", parent=styles["Normal"], fontSize=12, alignment=TA_CENTER
            )
            story.append(Paragraph(f"By: {author}", author_style))

        story.append(Spacer(1, 0.5 * inch))
        date_style = ParagraphStyle(
            "Date", parent=styles["Normal"], fontSize=10, alignment=TA_CENTER
        )
        story.append(Paragraph(datetime.now().strftime("%B %d, %Y"), date_style))

        story.append(PageBreak())

        # Table of Contents
        story.append(Paragraph("Table of Contents", heading_style))
        story.append(Spacer(1, 0.25 * inch))

        for i, section in enumerate(sections, 1):
            toc_style = ParagraphStyle(
                "TOC", parent=styles["Normal"], fontSize=11, leftIndent=20
            )
            story.append(
                Paragraph(f"{i}. {section.get('title', 'Untitled')}", toc_style)
            )

        story.append(PageBreak())

        # Sections
        for i, section in enumerate(sections, 1):
            # Section title
            story.append(
                Paragraph(f"{i}. {section.get('title', 'Untitled')}", heading_style)
            )

            # Section content
            content = section.get("content", "")
            if content:
                # Split content into paragraphs
                paragraphs = content.split("\n\n")
                for para in paragraphs:
                    if para.strip():
                        # Escape special characters
                        para = para.replace("&", "&amp;")
                        para = para.replace("<", "&lt;")
                        para = para.replace(">", "&gt;")
                        story.append(Paragraph(para, body_style))
            else:
                story.append(Paragraph("[No content]", body_style))

            story.append(Spacer(1, 0.25 * inch))

        # Build PDF
        doc.build(story)

        # Get PDF bytes
        pdf_bytes = buffer.getvalue()
        buffer.close()

        return pdf_bytes

    def export_to_docx(
        self,
        report_title: str,
        sections: list,
        author: str = "",
        metadata: Dict[str, Any] = None,
    ) -> bytes:
        """
        Export report to DOCX format.

        Args:
            report_title: Title of the report
            sections: List of sections with title and content
            author: Author name
            metadata: Additional metadata

        Returns:
            DOCX file as bytes
        """
        if not self.docx_available:
            raise ImportError(
                "python-docx is required for DOCX export. Install with: pip install python-docx"
            )

        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Create document
        doc = Document()

        # Set document properties
        core_props = doc.core_properties
        core_props.title = report_title
        if author:
            core_props.author = author

        # Title
        title_para = doc.add_heading(report_title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Author and date
        if author:
            author_para = doc.add_paragraph(f"By: {author}")
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        date_para = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Page break
        doc.add_page_break()

        # Table of Contents
        doc.add_heading("Table of Contents", level=1)
        for i, section in enumerate(sections, 1):
            doc.add_paragraph(
                f"{i}. {section.get('title', 'Untitled')}", style="List Number"
            )

        # Page break
        doc.add_page_break()

        # Sections
        for i, section in enumerate(sections, 1):
            # Section title
            doc.add_heading(f"{i}. {section.get('title', 'Untitled')}", level=1)

            # Section content
            content = section.get("content", "")
            if content:
                # Split content into paragraphs
                paragraphs = content.split("\n\n")
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para)
            else:
                doc.add_paragraph("[No content]")

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()

        return docx_bytes

    def get_export_filename(self, report_title: str, format: str) -> str:
        """
        Generate a safe filename for export.

        Args:
            report_title: Title of the report
            format: Export format (pdf or docx)

        Returns:
            Safe filename
        """
        import re

        # Remove unsafe characters
        safe_title = re.sub(r'[<>:"/\\|?*]', "", report_title)
        safe_title = safe_title.strip()[:50]  # Limit length

        if not safe_title:
            safe_title = "report"

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        return f"{safe_title}_{timestamp}.{format}"


# Singleton instance
export_service = ExportService()
